from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def extract_text(docx_path: Path) -> str:
    d = Document(str(docx_path))
    lines: list[str] = []

    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    for tb in d.tables:
        for row in tb.rows:
            cells: list[str] = []
            for c in row.cells:
                ct = (c.text or "").strip()
                if ct:
                    cells.append(ct)
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip() + "\n"


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: extract_docx_text.py <input.docx> <output.txt>", file=sys.stderr)
        return 2

    inp = Path(argv[1])
    out = Path(argv[2])
    out.parent.mkdir(parents=True, exist_ok=True)

    text = extract_text(inp)
    out.write_text(text, encoding="utf-8")
    print(str(out))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))

